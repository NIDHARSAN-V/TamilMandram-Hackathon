import os
import uuid
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional
from fastapi import Body
import whisper
from transcribe_utils import ensure_dir, normalize_to_wav, reduce_noise
from docx_utils import build_docx_from_segments
import requests
import openai
from io import BytesIO
from docx import Document
from fastapi.responses import StreamingResponse






GROQ_API_URL = "https://api.groq.com/v1"  # example endpoint
GROQ_API_KEY = "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"

client = openai.OpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1"
)

# Optional pyannote imports
USE_PYANNOTE = True
try:
    from pyannote.audio import Pipeline
    _pyannote_pipeline = None
except Exception:
    USE_PYANNOTE = False
    _pyannote_pipeline = None

UPLOAD_DIR = "/tmp/tamil_transcribe"
ensure_dir(UPLOAD_DIR)

app = FastAPI(title="Tamil Audio→Docx Transcriber")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

# Load Whisper model (env var WHISPER_MODEL)
MODEL_NAME = os.getenv("WHISPER_MODEL", "medium")
model = whisper.load_model(MODEL_NAME)

def get_pyannote_pipeline():
    global _pyannote_pipeline
    if not USE_PYANNOTE:
        return None
    if _pyannote_pipeline is None:
        hf_token = os.getenv("HUGGINGFACE_TOKEN")
        if not hf_token:
            raise RuntimeError("HUGGINGFACE_TOKEN required for pyannote diarization.")
        _pyannote_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization", use_auth_token=hf_token)
    return _pyannote_pipeline

def run_whisper(wav_path, language="ta"):
    # returns whisper result dict (with 'segments')
    result = model.transcribe(wav_path, language=language)
    return result

def run_pyannote(wav_path):
    pipeline = get_pyannote_pipeline()
    if pipeline is None:
        return None
    diar = pipeline(wav_path)
    diar_list = []
    for turn, _, speaker in diar.itertracks(yield_label=True):
        diar_list.append({"start": float(turn.start), "end": float(turn.end), "speaker": str(speaker)})
    return diar_list

def merge(whisper_segments, diarization_list=None):
    """
    Merge whisper segments with diarization list (if present) and produce a list of
    dictionaries: {speaker, start, end, text} with Tamil speaker names.
    """
    merged = []
    speaker_map = {}
    for seg in whisper_segments:
        s_start = seg.get("start", 0.0)
        s_end = seg.get("end", 0.0)
        text = seg.get("text", "").strip()
        label = None
        if diarization_list:
            # find diarization entry with max overlap
            overlaps = []
            for d in diarization_list:
                overlap = max(0.0, min(s_end, d["end"]) - max(s_start, d["start"]))
                if overlap > 0:
                    overlaps.append((overlap, d["speaker"]))
            if overlaps:
                overlaps.sort(reverse=True)
                label = overlaps[0][1]
        if label is None:
            # fallback heuristic: bucket by 30s windows
            bucket = int(s_start // 30) + 1
            label = f"Speaker_{bucket}"
        if label not in speaker_map:
            speaker_map[label] = f"பேச்சாளர் {len(speaker_map)+1}"
        tamil_label = speaker_map[label]
        merged.append({"speaker": tamil_label, "start": s_start, "end": s_end, "text": text})
    return merged

@app.post("/api/transcribe_text")
async def transcribe_text(
    audio: UploadFile = File(...),
    do_denoise: Optional[bool] = Form(True),
    do_diarize: Optional[bool] = Form(False)
):
    """
    Returns structured JSON segments (speaker, start, end, text) for in-browser editing.
    """
    uid = str(uuid.uuid4())
    raw_ext = Path(audio.filename).suffix or ".webm"
    raw_path = f"{UPLOAD_DIR}/{uid}_raw{raw_ext}"
    with open(raw_path, "wb") as f:
        f.write(await audio.read())

    wav_path = f"{UPLOAD_DIR}/{uid}.wav"
    normalize_to_wav(raw_path, wav_path, target_sr=16000)

    processed = wav_path
    if do_denoise:
        denoised = f"{UPLOAD_DIR}/{uid}_denoised.wav"
        try:
            reduce_noise(wav_path, denoised)
            processed = denoised
        except Exception:
            processed = wav_path

    # Whisper transcription
    whisper_result = run_whisper(processed, language="ta")
    segments = whisper_result.get("segments", [])

    diarization_list = None
    if do_diarize:
        try:
            diarization_list = run_pyannote(processed)
        except Exception as e:
            diarization_list = None
            
            

    merged = merge(segments, diarization_list)
    print(f"Transcription completed: {len(merged)} segments.")
    return JSONResponse({"segments": merged})


@app.post("/api/make_docx")
async def make_docx(segments: dict = None):
    """
    Accepts JSON body: { segments: [{speaker, start, end, text}, ...] }
    Returns .docx file.
    """
    if not segments or "segments" not in segments:
        raise HTTPException(status_code=400, detail="Missing 'segments' in request body")
    uid = str(uuid.uuid4())
    doc_path = f"{UPLOAD_DIR}/{uid}.docx"
    build_docx_from_segments(segments["segments"], doc_path, include_timestamps=True)
    return FileResponse(doc_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename="tamil_transcription.docx")
    


def call_groq_model(prompt: str, model: str = "llama-3.3-70b-versatile") -> str:
    """
    Calls Groq OpenAI-compatible API with a prompt.
    Returns the generated Tamil text.
    """
    try:
        response = client.responses.create(
            model=model,
            input=prompt
        )
        # Extract output text
        if hasattr(response, "output_text") and response.output_text:
            return response.output_text
        else:
            # Fallback: get first text chunk
            return response.output[0].content[0].text
    except Exception as e:
        print("Groq API error:", str(e))
        return f"Error calling Groq API: {str(e)}"


def call_groq_model(prompt: str, model: str = "llama-3.3-70b-versatile") -> str:
    """
    Calls Groq OpenAI-compatible API using chat completion endpoint.
    Returns Tamil text output.
    """
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1024
        )
        # Extract output text
        return response.choices[0].message.content
    except Exception as e:
        print("Groq API error:", str(e))
        return f"Error calling Groq API: {str(e)}"


@app.post("/api/min_meet")
async def generate_meeting_notes(segments: dict = Body(...)):
    if not segments or "segments" not in segments:
        raise HTTPException(status_code=400, detail="Missing 'segments'")
    
    valid_segments = []
    for s in segments["segments"]:
        speaker = s.get("speaker", "பேச்சாளர் ?")
        text = s.get("text", "").strip()
        if text:
            valid_segments.append(f"{speaker}: {text}")
    
    if not valid_segments:
        raise HTTPException(status_code=400, detail="No valid segments with text found")
    
    full_text = "\n".join(valid_segments)
    prompt = f"இந்த உரையை தமிழில் சுருக்கமாகக் குறிப்புகள் (bullet points) வடிவில் உருவாக்கவும்:\n{full_text}"
    
    notes = call_groq_model(prompt)
    return {"notes": notes}


@app.post("/api/summary")
async def generate_summary(segments: dict = Body(...)):
    if not segments or "segments" not in segments:
        raise HTTPException(status_code=400, detail="Missing 'segments'")
    
    valid_segments = []
    for s in segments["segments"]:
        speaker = s.get("speaker", "பேச்சாளர் ?")
        text = s.get("text", "").strip()
        if text:
            valid_segments.append(f"{speaker}: {text}")
    
    if not valid_segments:
        raise HTTPException(status_code=400, detail="No valid segments with text found")
    
    full_text = "\n".join(valid_segments)
    prompt = f"இந்த உரையை தமிழில் சுருக்கமாக விவரிக்கவும்:\n{full_text}"
    
    summary = call_groq_model(prompt)
    return {"summary": summary}




def build_docx_from_text(text: str, title: str = "Document") -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

@app.post("/api/download_notes_docx")
async def download_notes_docx(segments: dict = Body(...)):
    if not segments or "segments" not in segments:
        raise HTTPException(status_code=400, detail="Missing 'segments'")
    
    valid_segments = []
    for s in segments["segments"]:
        speaker = s.get("speaker", "பேச்சாளர் ?")
        text = s.get("text", "").strip()
        if text:
            valid_segments.append(f"{speaker}: {text}")
    full_text = "\n".join(valid_segments)

    # Generate meeting notes via Groq
    prompt = f"இந்த உரையை தமிழில் சுருக்கமாகக் குறிப்புகள் (bullet points) வடிவில் உருவாக்கவும்:\n{full_text}"
    notes = call_groq_model(prompt)

    docx_file = build_docx_from_text(notes, title="Meeting Notes")
    return StreamingResponse(docx_file,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=meeting_notes.docx"})


@app.post("/api/download_summary_docx")
async def download_summary_docx(segments: dict = Body(...)):
    if not segments or "segments" not in segments:
        raise HTTPException(status_code=400, detail="Missing 'segments'")
    
    valid_segments = []
    for s in segments["segments"]:
        speaker = s.get("speaker", "பேச்சாளர் ?")
        text = s.get("text", "").strip()
        if text:
            valid_segments.append(f"{speaker}: {text}")
    full_text = "\n".join(valid_segments)

    # Generate summary via Groq
    prompt = f"இந்த உரையை தமிழில் சுருக்கமாக விவரிக்கவும்:\n{full_text}"
    summary = call_groq_model(prompt)

    docx_file = build_docx_from_text(summary, title="Summary")
    return StreamingResponse(docx_file,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=summary.docx"})
