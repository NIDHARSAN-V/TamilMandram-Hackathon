from docx import Document

def build_docx_from_segments(segments, doc_path, title="தமிழ் உரை (Transcription)", include_timestamps=True):
    """
    segments: list of {speaker, start, end, text}
    """
    doc = Document()
    doc.add_heading(title, level=1)
    for seg in segments:
        header = f"{seg.get('speaker','பேச்சாளர்')} "
        if include_timestamps:
            header += f"[{seg.get('start',0):.2f}s - {seg.get('end',0):.2f}s]"
        p = doc.add_paragraph()
        p.add_run(header + "\n").bold = True
        p.add_run(seg.get('text','') + "\n\n")
    doc.save(doc_path)
    return doc_path
